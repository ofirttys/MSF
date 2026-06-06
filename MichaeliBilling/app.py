"""
MichaeliBilling — Dr. J. Michaeli, Mount Sinai Fertility
Billing dashboard: import eIVF XLS → review/edit → export report → save to DB
"""

import eel
import sqlite3
import pandas as pd
import json
import os
import sys
import logging
import tempfile
from datetime import datetime
from pathlib import Path
from billing_rules import assign_billing_codes

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE_DIR   = Path(sys.executable).parent if getattr(sys, 'frozen', False) else Path(__file__).parent
DB_PATH    = BASE_DIR / "db"      / "billing.db"
EXPORT_DIR = BASE_DIR / "exports"
LOG_DIR    = BASE_DIR / "logs"
WEB_DIR    = BASE_DIR / "web"

for d in (DB_PATH.parent, EXPORT_DIR, LOG_DIR):
    d.mkdir(parents=True, exist_ok=True)

# ── Logging ───────────────────────────────────────────────────────────────────
logging.basicConfig(
    filename=LOG_DIR / f"billing_{datetime.now():%Y%m%d}.log",
    level=logging.INFO,
    format="%(asctime)s  %(levelname)s  %(message)s",
)
log = logging.getLogger("MichaeliBilling")

# ── Reference data ────────────────────────────────────────────────────────────
# Billing codes organised by clinical group.
# Each code: {"desc": str, "fee": float, "group": str}
BILLING_CODES = {
    # ── Consultations ──────────────────────────────────────────────────────────
    "A205": {"desc": "Consultation*",                  "fee": 134.25, "group": "Consultations"},
    "A935": {"desc": "Special surgical consultation",  "fee": 194.65, "group": "Consultations"},
    "A206": {"desc": "Repeat consultation",            "fee":  71.45, "group": "Consultations"},
    "A203": {"desc": "Specific assessment",            "fee":  62.65, "group": "Consultations"},
    "A204": {"desc": "Partial assessment",             "fee":  40.50, "group": "Consultations"},
    # ── Counselling ────────────────────────────────────────────────────────────
    "K013": {"desc": "Individual Counselling x3",      "fee":  80.00, "group": "Counselling"},
    "K033": {"desc": "Individual Counselling",         "fee":  56.30, "group": "Counselling"},
    "K040": {"desc": "Group counselling x3",           "fee":  80.00, "group": "Counselling"},
    "K041": {"desc": "Group counselling",              "fee":  56.30, "group": "Counselling"},
    # ── Virtual / Phone ────────────────────────────────────────────────────────
    "A101": {"desc": "Limited Virtual Care by Video",     "fee": 20.00, "group": "Virtual / Phone"},
    "A102": {"desc": "Limited Virtual Care by Telephone", "fee": 15.00, "group": "Virtual / Phone"},
    "K300": {"desc": "Virtual appointment",               "fee":  0.00, "group": "Virtual / Phone"},
    "K301": {"desc": "Phone call",                        "fee":  0.00, "group": "Virtual / Phone"},
    # ── Ultrasound ─────────────────────────────────────────────────────────────
    "J164C": {"desc": "Follicle monitoring",                    "fee": 0.00, "group": "Ultrasound"},
    "J164B": {"desc": "Follicle monitoring",                    "fee": 0.00, "group": "Ultrasound"},
    "J476B": {"desc": "Transvaginal sonohysterography (Echovist)", "fee": 0.00, "group": "Ultrasound"},
    "J476C": {"desc": "Transvaginal sonohysterography (Echovist)", "fee": 0.00, "group": "Ultrasound"},
    "J165B": {"desc": "Transvaginal sonohysterography",         "fee": 0.00, "group": "Ultrasound"},
    "J165C": {"desc": "Transvaginal sonohysterography",         "fee": 0.00, "group": "Ultrasound"},
    "J138B": {"desc": "Intracavitary ultrasound",               "fee": 0.00, "group": "Ultrasound"},
    "J138C": {"desc": "Intracavitary ultrasound",               "fee": 0.00, "group": "Ultrasound"},
    "J149":  {"desc": "Ultrasonic guidance of biopsy/aspiration/amniocentesis", "fee": 0.00, "group": "Ultrasound"},
    # ── Procedures ─────────────────────────────────────────────────────────────
    "G399A": {"desc": "Transvaginal sonohysterography, intro of catheter", "fee": 0.00, "group": "Procedures"},
    "J008A": {"desc": "Hysterosalpingogram",                               "fee": 0.00, "group": "Procedures"},
    "E861":  {"desc": "Paracervical block",                                "fee": 0.00, "group": "Procedures"},
    "Z770":  {"desc": "Endometrial sampling",                              "fee": 0.00, "group": "Procedures"},
    "Z582":  {"desc": "Hysteroscopy (diagnostic)",                         "fee": 0.00, "group": "Procedures"},
    "Z583":  {"desc": "Hysteroscopy (with uterine biopsy)",                "fee": 0.00, "group": "Procedures"},
    "Z587":  {"desc": "Hysteroscopy (with resection of polyps/fibroids)",  "fee": 0.00, "group": "Procedures"},
    "Z585":  {"desc": "Hysteroscopy (with cannulization of tubes)",        "fee": 0.00, "group": "Procedures"},
    "S756":  {"desc": "Missed abortion / evacuation of molar pregnancy",   "fee": 0.00, "group": "Procedures"},
}

# Group order for display
BILLING_CODE_GROUPS = ["Consultations", "Counselling", "Virtual / Phone", "Ultrasound", "Procedures"]

DX_CODES = {
    "628": "Other Disorders of Female Genital Tract: Infertility",
    "606": "Diseases of Male Genital Organs: Male infertility, oligospermia, azoospermia",
    "626": "Other Disorders of Female Genital Tract: Disorders of menstruation",
    "625": "Other Disorders of Female Genital Tract: Dyspareunia, dysmenorrhea, premenstrual tension, stress incontinence",
    "895": "Family Planning: Family planning, contraceptive advice, advice on surgical contraception or abortion",
    "758": "Chromosomal anomalies (e.g., Down's syndrome, other autosomal anomalies, Klinefelter's syndrome, Turner's syndrome)",
    "632": "Missed abortion",
    "633": "Ectopic pregnancy",
    "634": "Incomplete abortion, complete abortion",
    "635": "Therapeutic abortion",
    "640": "Threatened abortion, haemorrhage in early pregnancy",
    "650": "Normal delivery, uncomplicated pregnancy",
}

# ── DB setup ──────────────────────────────────────────────────────────────────
def init_db():
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.executescript("""
        CREATE TABLE IF NOT EXISTS sessions (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            session_date  TEXT NOT NULL,
            imported_at   TEXT NOT NULL,
            source_file   TEXT,
            submitted     INTEGER DEFAULT 0,
            submitted_at  TEXT
        );

        CREATE TABLE IF NOT EXISTS encounters (
            id                   INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id           INTEGER NOT NULL REFERENCES sessions(id),
            patient_id           TEXT,
            partner_id           TEXT,
            patient_name         TEXT NOT NULL,
            health_card          TEXT,
            visit_type           TEXT,
            facility             TEXT,
            status               TEXT,
            encounter_date       TEXT NOT NULL,
            start_time           TEXT,
            end_time             TEXT,
            duration_min         INTEGER,
            referring_md         TEXT,
            referring_md_license TEXT,
            schedule_notes       TEXT,
            last_encounter_date  TEXT,
            last_encounter_type  TEXT,
            months_since_last    INTEGER,
            provider_enc_count   INTEGER,
            billing_codes        TEXT,
            dx_codes             TEXT,
            sex                  TEXT,
            notes                TEXT,
            locked               INTEGER DEFAULT 0,
            flag_level           TEXT DEFAULT "",
            flag_messages        TEXT DEFAULT "[]",
            included             INTEGER DEFAULT 1,
            md_copied            INTEGER DEFAULT 0
        );

        CREATE TABLE IF NOT EXISTS patient_records (
            id            INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_id    TEXT,
            patient_name  TEXT NOT NULL,
            record_date   TEXT NOT NULL,
            billing_codes TEXT NOT NULL,
            dx_codes      TEXT NOT NULL,
            source        TEXT,
            source_ref    TEXT,
            created_at    TEXT NOT NULL
        );

        CREATE INDEX IF NOT EXISTS idx_pr_patient_id   ON patient_records(patient_id);
        CREATE INDEX IF NOT EXISTS idx_pr_patient_name ON patient_records(patient_name COLLATE NOCASE);
        CREATE INDEX IF NOT EXISTS idx_pr_record_date  ON patient_records(record_date);

        CREATE TABLE IF NOT EXISTS historical_imports (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            imported_at TEXT NOT NULL,
            source_file TEXT,
            row_count   INTEGER
        );
    """)
    con.commit()
    con.close()
    log.info("Database initialised at %s", DB_PATH)

# ── Helpers ───────────────────────────────────────────────────────────────────
def parse_sex(raw: str) -> str:
    """Normalise Ptn_Sex column value → 'F' | 'M' | 'U'."""
    v = str(raw or "").strip().upper()
    if v in ("F", "FEMALE"):
        return "F"
    if v in ("M", "MALE"):
        return "M"
    return "U"

def default_dx(sex: str) -> list:
    if sex == "F":
        return ["628"]
    if sex == "M":
        return ["606"]
    return ["628"]

def referring_license(row) -> str:
    for col in ("State_License_No", "Upin"):
        val = str(row.get(col, "") or "").strip()
        if val and val.lower() not in ("nan", ""):
            if "." in val:
                try:
                    val = str(int(float(val)))
                except ValueError:
                    pass
            return val
    return ""

def db_con():
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row
    return con

def safe_int(val):
    try:
        return int(val) if pd.notna(val) else None
    except (ValueError, TypeError):
        return None

# ── Eel API ───────────────────────────────────────────────────────────────────

@eel.expose
def get_reference_data():
    return {
        "billing_codes":        BILLING_CODES,
        "billing_code_groups":  BILLING_CODE_GROUPS,
        "dx_codes":             DX_CODES,
    }

@eel.expose
def import_xls(file_path: str):
    """Parse XLS/XLSX daily billing export from a filesystem path."""
    try:
        ext    = Path(file_path).suffix.lower()
        engine = "xlrd" if ext == ".xls" else "openpyxl"
        df     = pd.read_excel(file_path, engine=engine)

        required = {"Visit_Type", "Schedule_Date", "From_Time", "To_Time",
                    "Scheduled_EntityName", "Patient_ID"}
        missing  = required - set(df.columns)
        if missing:
            return {"ok": False, "error": f"Missing columns: {', '.join(sorted(missing))}"}

        has_sex_col = "Ptn_Sex" in df.columns
        encounters  = []

        for _, row in df.iterrows():
            row = row.to_dict()

            # Sex: prefer explicit Ptn_Sex column, fall back to Visit_Type inference
            if has_sex_col and pd.notna(row.get("Ptn_Sex")):
                sex = parse_sex(row["Ptn_Sex"])
            else:
                vt  = str(row.get("Visit_Type", "")).upper()
                sex = "F" if ("FEMALE" in vt or "NPF" in vt) else "M" if "MALE" in vt else "U"

            partner_raw = row.get("Ptn_Partner_Id")
            partner_id  = str(int(partner_raw)) if pd.notna(partner_raw) else ""

            enc = {
                "patient_id":           str(row.get("Patient_ID", "") or ""),
                "partner_id":           partner_id,
                "patient_name":         str(row.get("Scheduled_EntityName", "")).strip(),
                "health_card":          str(row.get("Ptn_SSN", "") or "").strip(),
                "visit_type":           str(row.get("Visit_Type", "") or "").strip(),
                "facility":             str(row.get("Facility_Name", "") or "").strip(),
                "status":               str(row.get("Status", "") or "").strip(),
                "encounter_date":       str(row.get("Schedule_Date", ""))[:10],
                "start_time":           str(row.get("From_Time", "") or "").strip(),
                "end_time":             str(row.get("To_Time", "") or "").strip(),
                "duration_min":         safe_int(row.get("Duration")),
                "referring_md":         str(row.get("Ptn_Referring_MD_Name", "") or "").strip(),
                "referring_md_license": referring_license(row),
                "schedule_notes":       str(row.get("Schedule_Notes", "") or "").strip(),
                "last_encounter_date":  str(row.get("LatestNPEncounterDate", ""))[:10]
                                        if pd.notna(row.get("LatestNPEncounterDate")) else "",
                "last_encounter_type":  str(row.get("LatestNPEncounterVisitType", "") or "").strip(),
                "months_since_last":    safe_int(row.get("MonthsSinceLastEncounter")),
                "provider_enc_count":   safe_int(row.get("ProviderEncounterCount")),
                "sex":                  sex,
                "billing_codes":        [],
                "dx_codes":             default_dx(sex),
                "notes":                "",
                "flag_level":           "",
                "flag_messages":        [],
                "included":             True,
                "md_copied":            False,
            }
            encounters.append(enc)

        # Run billing rules engine
        encounters = assign_billing_codes(encounters, DB_PATH)

        session_date = encounters[0]["encounter_date"] if encounters else ""
        log.info("Imported %d encounters from %s", len(encounters), file_path)
        return {
            "ok":           True,
            "session_date": session_date,
            "source_file":  Path(file_path).name,
            "has_sex_col":  has_sex_col,
            "encounters":   encounters,
        }

    except Exception as e:
        log.exception("import_xls failed")
        return {"ok": False, "error": str(e)}

@eel.expose
def import_xls_bytes(byte_array: list, filename: str):
    """
    Receive raw file bytes from the browser (drag-and-drop or file picker),
    write to a temp file, parse it, then delete the temp file.
    This works around the browser security restriction that prevents JavaScript
    from accessing the full filesystem path of a dropped/selected file.
    """
    try:
        ext      = Path(filename).suffix.lower()
        tmp_path = Path(tempfile.mktemp(suffix=ext))
        tmp_path.write_bytes(bytes(byte_array))

        result = import_xls(str(tmp_path))

        tmp_path.unlink(missing_ok=True)

        # Restore the original filename (import_xls would have used the temp name)
        if result.get("ok"):
            result["source_file"] = filename

        return result

    except Exception as e:
        log.exception("import_xls_bytes failed")
        return {"ok": False, "error": str(e)}

@eel.expose
def save_session(session_date: str, source_file: str, encounters: list):
    """Commit a reviewed session to the database (explicit user action only)."""
    try:
        con = db_con()
        cur = con.cursor()
        cur.execute(
            "INSERT INTO sessions (session_date, imported_at, source_file) VALUES (?,?,?)",
            (session_date, datetime.now().isoformat(), source_file)
        )
        session_id = cur.lastrowid

        for e in encounters:
            cur.execute("""
                INSERT INTO encounters
                  (session_id, patient_id, partner_id, patient_name, health_card,
                   visit_type, facility, status, encounter_date, start_time, end_time,
                   duration_min, referring_md, referring_md_license, schedule_notes,
                   last_encounter_date, last_encounter_type, months_since_last,
                   provider_enc_count, billing_codes, dx_codes, sex, notes,
                   flag_level, flag_messages, included, md_copied)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """, (
                session_id,
                e.get("patient_id",""),         e.get("partner_id",""),
                e.get("patient_name",""),        e.get("health_card",""),
                e.get("visit_type",""),          e.get("facility",""),
                e.get("status",""),              e.get("encounter_date",""),
                e.get("start_time",""),          e.get("end_time",""),
                e.get("duration_min"),
                e.get("referring_md",""),        e.get("referring_md_license",""),
                e.get("schedule_notes",""),
                e.get("last_encounter_date",""), e.get("last_encounter_type",""),
                e.get("months_since_last"),      e.get("provider_enc_count"),
                json.dumps(e.get("billing_codes", [])),
                json.dumps(e.get("dx_codes", [])),
                e.get("sex",""),                 e.get("notes",""),
                e.get("flag_level",""),
                json.dumps(e.get("flag_messages",[])),
                1 if e.get("included", True) else 0,
                1 if e.get("md_copied", False) else 0,
            ))

        # Also write to patient_records (the clean billing ledger)
        for e in encounters:
            if not e.get("patient_name","").strip():
                continue
            cur.execute("""
                INSERT INTO patient_records
                  (patient_id, patient_name, record_date, billing_codes, dx_codes,
                   source, source_ref, created_at)
                VALUES (?,?,?,?,?,?,?,?)
            """, (
                e.get("patient_id",""),
                e.get("patient_name","").strip(),
                e.get("encounter_date",""),
                json.dumps(e.get("billing_codes",[])),
                json.dumps(e.get("dx_codes",[])),
                "session",
                str(session_id),
                datetime.now().isoformat(),
            ))

        con.commit()
        con.close()
        log.info("Saved session #%d (%s, %d encounters)", session_id, session_date, len(encounters))
        return {"ok": True, "session_id": session_id}

    except Exception as e:
        log.exception("save_session failed")
        return {"ok": False, "error": str(e)}

@eel.expose
def get_sessions():
    con  = db_con()
    rows = con.execute("""
        SELECT s.id, s.session_date, s.imported_at, s.source_file,
               s.submitted, s.submitted_at,
               COUNT(e.id) AS encounter_count
        FROM sessions s
        LEFT JOIN encounters e ON e.session_id = s.id
        GROUP BY s.id
        ORDER BY s.session_date DESC
    """).fetchall()
    con.close()
    return [dict(r) for r in rows]

@eel.expose
def get_session_encounters(session_id: int):
    con  = db_con()
    rows = con.execute(
        "SELECT * FROM encounters WHERE session_id=? ORDER BY start_time",
        (session_id,)
    ).fetchall()
    con.close()
    result = []
    for r in rows:
        d = dict(r)
        d["billing_codes"]  = json.loads(d["billing_codes"]  or "[]")
        d["dx_codes"]       = json.loads(d["dx_codes"]       or "[]")
        d["flag_messages"]  = json.loads(d.get("flag_messages") or "[]")
        d["included"]       = bool(d.get("included", 1))
        d["md_copied"]      = bool(d.get("md_copied", 0))
        result.append(d)
    return result

@eel.expose
def update_encounter(encounter_id: int, billing_codes: list, dx_codes: list,
                     notes: str, start_time: str = "", end_time: str = "",
                     health_card: str = "", referring_md: str = "",
                     referring_md_license: str = "", included: bool = True):
    try:
        con = db_con()
        con.execute(
            """UPDATE encounters
               SET billing_codes=?, dx_codes=?, notes=?,
                   start_time=?, end_time=?,
                   health_card=?, referring_md=?, referring_md_license=?,
                   included=?
               WHERE id=?""",
            (json.dumps(billing_codes), json.dumps(dx_codes), notes,
             start_time, end_time,
             health_card, referring_md, referring_md_license,
             1 if included else 0,
             encounter_id)
        )
        con.commit()
        con.close()
        return {"ok": True}
    except Exception as e:
        return {"ok": False, "error": str(e)}

@eel.expose
def mark_submitted(session_id: int):
    try:
        con = db_con()
        con.execute(
            "UPDATE sessions SET submitted=1, submitted_at=? WHERE id=?",
            (datetime.now().isoformat(), session_id)
        )
        con.commit()
        con.close()
        log.info("Session #%d marked submitted", session_id)
        return {"ok": True}
    except Exception as e:
        return {"ok": False, "error": str(e)}

@eel.expose
def export_report(session_id: int, encounters: list, fmt: str):
    """Generate billing report. fmt: 'xlsx' | 'docx' | 'pdf'"""
    try:
        rows = []
        for e in encounters:
            # Respect the included checkbox — skip excluded rows
            if not e.get("included", True):
                continue
            codes     = e.get("billing_codes", [])
            dxs       = e.get("dx_codes", [])
            # Codes and Dx as numbers only, semicolon-separated
            dx_str    = "; ".join(dxs)
            code_str  = "; ".join(codes)
            fee_total = sum(BILLING_CODES.get(c, {}).get("fee", 0) for c in codes)

            def clean(v):
                """Return empty string instead of nan/None."""
                if v is None: return ""
                s = str(v).strip()
                return "" if s.lower() == "nan" else s

            rows.append({
                "Date":                clean(e.get("encounter_date")),
                "Start Time":          clean(e.get("start_time")),
                "End Time":            clean(e.get("end_time")),
                "Patient Name":        clean(e.get("patient_name")),
                "Health Card":         clean(e.get("health_card")),
                "Sex":                 clean(e.get("sex")),
                "Dx":                  dx_str,
                "Billing Code(s)":     code_str,
                "Referring Physician": clean(e.get("referring_md")),
                "Billing#":            clean(e.get("referring_md_license")),
                "Notes":               clean(e.get("notes")),
            })

        df           = pd.DataFrame(rows)
        session_date = rows[0]["Date"] if rows else datetime.now().strftime("%Y-%m-%d")

        # Running report ID: count existing files for this date across all formats
        existing  = list(EXPORT_DIR.glob(f"DrMichaeli-BillingReport_{session_date}_*"))
        report_id = len(existing) + 1
        base_name = f"DrMichaeli-BillingReport_{session_date}_{report_id}"

        # ── XLSX ──────────────────────────────────────────────────────────────
        if fmt == "xlsx":
            from openpyxl.styles import Font, PatternFill, Alignment
            from openpyxl.utils import get_column_letter

            out = EXPORT_DIR / f"{base_name}.xlsx"
            with pd.ExcelWriter(out, engine="openpyxl") as writer:
                # Write two header rows then data
                df.to_excel(writer, index=False, sheet_name="Billing Report", startrow=3)
                ws = writer.sheets["Billing Report"]

                # Row 1: title, Row 2: date, Row 3: blank
                ws.cell(1, 1).value = "Dr. Michaeli - Billing Report"
                ws.cell(1, 1).font  = Font(bold=True, size=13, color="1B3A5C")
                ws.cell(2, 1).value = f"Date: {session_date}"
                ws.cell(2, 1).font  = Font(size=10, color="4A6A82")

                # Style column header row (row 4)
                header_fill = PatternFill("solid", fgColor="1B3A5C")
                header_font = Font(bold=True, color="FFFFFF", size=10)
                for cell in ws[4]:
                    cell.fill      = header_fill
                    cell.font      = header_font
                    cell.alignment = Alignment(horizontal="center", wrap_text=True)

                # Auto-size columns
                for col_idx, col in enumerate(ws.columns, 1):
                    max_len = max((len(str(c.value or "")) for c in col), default=8)
                    ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 55)

                # Alternate row shading from row 5 onward
                from openpyxl.styles import PatternFill as PF
                alt_fill = PF("solid", fgColor="EEF2F7")
                for row_idx, row in enumerate(ws.iter_rows(min_row=5), 5):
                    if row_idx % 2 == 0:
                        for cell in row:
                            cell.fill = alt_fill

            return {"ok": True, "path": str(out)}

        # ── DOCX ──────────────────────────────────────────────────────────────
        elif fmt == "docx":
            from docx import Document
            from docx.shared import Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            doc = Document()

            for section in doc.sections:
                section.left_margin   = Cm(1.5)
                section.right_margin  = Cm(1.5)
                section.top_margin    = Cm(1.8)
                section.bottom_margin = Cm(1.8)

            # Title
            title     = doc.add_heading("", 0)
            title_run = title.add_run("Dr. Michaeli - Billing Report")
            title_run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

            # Date line
            date_p   = doc.add_paragraph()
            date_run = date_p.add_run(f"Date: {session_date}")
            date_run.font.size      = Pt(11)
            date_run.font.color.rgb = RGBColor(0x4A, 0x6A, 0x82)

            doc.add_paragraph("")

            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = "Table Grid"

            for i, col_name in enumerate(df.columns):
                cell      = table.rows[0].cells[i]
                cell.text = col_name
                run       = cell.paragraphs[0].runs[0]
                run.bold  = True
                run.font.size      = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                tc_pr = cell._tc.get_or_add_tcPr()
                shd   = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "1B3A5C")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:val"), "clear")
                tc_pr.append(shd)

            for row_data in df.itertuples(index=False):
                cells = table.add_row().cells
                for i, val in enumerate(row_data):
                    cells[i].text = str(val)
                    cells[i].paragraphs[0].runs[0].font.size = Pt(8)

            out = EXPORT_DIR / f"{base_name}.docx"
            doc.save(out)
            return {"ok": True, "path": str(out)}

        # ── PDF ───────────────────────────────────────────────────────────────
        elif fmt == "pdf":
            from reportlab.lib.pagesizes import landscape, A4
            from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                            Paragraph, Spacer)
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm

            out     = EXPORT_DIR / f"{base_name}.pdf"
            doc_pdf = SimpleDocTemplate(
                str(out), pagesize=landscape(A4),
                leftMargin=1*cm, rightMargin=1*cm,
                topMargin=1.5*cm, bottomMargin=1.5*cm
            )

            styles      = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "MBTitle", parent=styles["Heading1"],
                fontSize=14, textColor=colors.HexColor("#1B3A5C"),
                spaceAfter=4,
            )
            date_style = ParagraphStyle(
                "MBDate", parent=styles["Normal"],
                fontSize=10, textColor=colors.HexColor("#4A6A82"),
                spaceAfter=12,
            )

            story = [
                Paragraph("Dr. Michaeli - Billing Report", title_style),
                Paragraph(f"Date: {session_date}", date_style),
            ]

            header = list(df.columns)
            data   = [header] + [[str(v) for v in row] for _, row in df.iterrows()]

            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND",     (0, 0), (-1, 0), colors.HexColor("#1B3A5C")),
                ("TEXTCOLOR",      (0, 0), (-1, 0), colors.white),
                ("FONTNAME",       (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE",       (0, 0), (-1, -1), 7),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                    [colors.white, colors.HexColor("#EEF2F7")]),
                ("GRID",           (0, 0), (-1, -1), 0.3, colors.HexColor("#CCCCCC")),
                ("VALIGN",         (0, 0), (-1, -1), "TOP"),
                ("PADDING",        (0, 0), (-1, -1), 3),
            ]))
            story.append(t)

            doc_pdf.build(story)
            return {"ok": True, "path": str(out)}

        return {"ok": False, "error": f"Unknown format: {fmt}"}

    except Exception as e:
        log.exception("export_report failed")
        import traceback
        return {"ok": False, "error": traceback.format_exc()}

@eel.expose
def import_historical_csv_bytes(byte_array: list, filename: str):
    """
    Import historical billing records from a CSV file sent as bytes from the browser.

    Expected CSV columns (case-insensitive, any order):
        patient_id   — optional
        patient_name — required
        date         — YYYY-MM-DD
        billing_codes — semicolon-separated codes e.g. "A205;K013"
        dx_codes      — semicolon-separated codes e.g. "628;606"
    """
    import tempfile, csv, io
    try:
        raw = bytes(byte_array).decode("utf-8-sig")   # handle BOM if present
        reader = csv.DictReader(io.StringIO(raw))

        # Normalise column names to lower-case stripped
        def norm(d):
            return {k.strip().lower(): v.strip() for k, v in d.items()}

        con = db_con()
        cur = con.cursor()

        # Log the import
        cur.execute(
            "INSERT INTO historical_imports (imported_at, source_file, row_count) VALUES (?,?,?)",
            (datetime.now().isoformat(), filename, 0)
        )
        import_id = cur.lastrowid

        inserted = 0
        skipped  = 0
        errors   = []

        for i, raw_row in enumerate(reader, start=2):   # row 2 = first data row
            row = norm(raw_row)

            name = row.get("patient_name","").strip()
            date = row.get("date","").strip()

            if not name or not date:
                skipped += 1
                errors.append(f"Row {i}: missing patient_name or date — skipped")
                continue

            # Parse codes — accept semicolons or commas as separators
            def parse_codes(val):
                if not val: return []
                return [c.strip() for c in val.replace(",",";").split(";") if c.strip()]

            billing = parse_codes(row.get("billing_codes",""))
            dx      = parse_codes(row.get("dx_codes",""))

            if not billing:
                skipped += 1
                errors.append(f"Row {i}: {name} — no billing codes, skipped")
                continue

            cur.execute("""
                INSERT INTO patient_records
                  (patient_id, patient_name, record_date, billing_codes, dx_codes,
                   source, source_ref, created_at)
                VALUES (?,?,?,?,?,?,?,?)
            """, (
                row.get("patient_id",""),
                name, date,
                json.dumps(billing),
                json.dumps(dx),
                "csv_import",
                filename,
                datetime.now().isoformat(),
            ))
            inserted += 1

        # Update row count
        cur.execute("UPDATE historical_imports SET row_count=? WHERE id=?", (inserted, import_id))
        con.commit()
        con.close()

        log.info("CSV import: %d inserted, %d skipped from %s", inserted, skipped, filename)
        return {
            "ok": True,
            "inserted": inserted,
            "skipped":  skipped,
            "errors":   errors[:20],   # cap error list
        }

    except Exception as e:
        log.exception("import_historical_csv_bytes failed")
        return {"ok": False, "error": str(e)}


@eel.expose
def search_patient_records(query: str = "", date: str = "", page: int = 1, page_size: int = 50):
    """
    Search patient_records.
    query  — matches patient_name or patient_id (partial, case-insensitive)
    date   — exact date filter YYYY-MM-DD (optional)
    Returns paginated results + total count.
    """
    try:
        con  = db_con()
        q    = query.strip()
        d    = date.strip()

        where_parts = []
        params      = []

        if q:
            where_parts.append("(patient_name LIKE ? OR patient_id LIKE ?)")
            params += [f"%{q}%", f"%{q}%"]
        if d:
            where_parts.append("record_date = ?")
            params.append(d)

        where = ("WHERE " + " AND ".join(where_parts)) if where_parts else ""

        total = con.execute(
            f"SELECT COUNT(*) FROM patient_records {where}", params
        ).fetchone()[0]

        offset = (page - 1) * page_size
        rows   = con.execute(
            f"""SELECT * FROM patient_records {where}
                ORDER BY record_date DESC, patient_name ASC
                LIMIT ? OFFSET ?""",
            params + [page_size, offset]
        ).fetchall()
        con.close()

        result = []
        for r in rows:
            d2 = dict(r)
            d2["billing_codes"] = json.loads(d2["billing_codes"] or "[]")
            d2["dx_codes"]      = json.loads(d2["dx_codes"]      or "[]")
            result.append(d2)

        return {"ok": True, "records": result, "total": total, "page": page, "page_size": page_size}

    except Exception as e:
        log.exception("search_patient_records failed")
        return {"ok": False, "error": str(e), "records": [], "total": 0}


@eel.expose
def get_patient_history(patient_id: str = "", patient_name: str = ""):
    """Return all billing records for a specific patient (by ID or exact name)."""
    try:
        con = db_con()
        if patient_id:
            rows = con.execute(
                "SELECT * FROM patient_records WHERE patient_id=? ORDER BY record_date DESC",
                (patient_id,)
            ).fetchall()
        else:
            rows = con.execute(
                "SELECT * FROM patient_records WHERE patient_name LIKE ? ORDER BY record_date DESC",
                (f"%{patient_name}%",)
            ).fetchall()
        con.close()

        result = []
        for r in rows:
            d = dict(r)
            d["billing_codes"] = json.loads(d["billing_codes"] or "[]")
            d["dx_codes"]      = json.loads(d["dx_codes"]      or "[]")
            result.append(d)

        return {"ok": True, "records": result}

    except Exception as e:
        return {"ok": False, "error": str(e), "records": []}

# ── Launch ────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    init_db()
    eel.init(str(WEB_DIR))

    print("Starting MichaeliBilling...")
    try:
        print("Trying Microsoft Edge...")
        eel.start("index.html", size=(1440, 900), port=8765, mode="edge", block=True)
    except (OSError, Exception) as e:
        print(f"Edge not available: {e}")
        try:
            print("Trying Chrome...")
            eel.start("index.html", size=(1440, 900), port=8765, mode="chrome", block=True)
        except (OSError, Exception) as e2:
            print(f"Chrome not available: {e2}")
            print("Falling back to default browser...")
            eel.start("index.html", size=(1440, 900), port=8765, mode=None, block=True)
